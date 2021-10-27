#This work is licensed under a Creative Commons Attribution-NonCommercial 4.0 International License.
#created by Hadi Muhammed 2021
import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

flag = 1
count = 1
steps = 1
ch = 0
where = " INT. "
time = " - DAY"
title = "0"
action = "0"
charac = {}
index = 1
count = int(input("Enter the last title count : "))
while(flag):

    def ask_where():
        global ch,where,time,title,action
        ch = input("Enter '1' for EXT , '2' for INT : ")
        if(int(ch) == 1):
            where = " EXT. "
        elif(int(ch) == 2):
            where = " INT. "
        else:
            ask_where()

    def ask_title():
        global ch, where, time, title, action
        title = input("Enter a meaningful title : ")
        if(len(title) < 3):
            ask_title()
        if(len(title) > 70):
            ask_title()
        title = title.upper()
    ch = 0
    def ask_time():
        global ch, where, time, title, action
        ch = input("Enter a '1' for 'Day' , '2' for 'Night' , '3' for Customizing: ")
        if(int(ch) == 1):
            time = " - DAY"
        elif(int(ch) == 2):
            time = " - NIGHT"
        elif(int(ch) == 3):
            time = " - "
            string = input("Enter a custom string : ")
            time = time + string.upper()
        else:
            ask_time()

    def ask_action():
        global ch, where, time, title, action
        action = input("Enter the action here : ")
        if(len(action) < 4):
            ask_action()

    def save_title():
        script = docx.Document()

        try:
            script = docx.Document("file.docx")
            print("[Add next scene]")
        except:
            print("[New document created]")

        text = script.add_heading().add_run(str(count)+where+title+time)
        text.font.size = Pt(12)
        text.font.name = 'Courier'
        text = script.add_paragraph().add_run(action)
        text.font.size = Pt(12)
        text.font.name = 'Courier'
        script.save('file.docx')

    def save_char(characters):
        script = docx.Document("file.docx")
        for key in characters:



            text = script.add_paragraph()
            run = text.add_run(str(characters[key][0]) + "\n" + str(characters[key][1]) + "\n" + str(characters[key][2]))
            font = run.font
            font.name = "Courier"
            font.size = Pt(12)
            text.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER






        script.save('file.docx')

    if(steps == 1):
        print("Your story settings :- ")
        if("INT" in where):
            print("[1] Place is - INTERIOR")
        if("EXT" in where):
            print("[1] Place is - EXTERIOR")
        if(title == "0"):
            print("[2] Title -  Not Set")
        else:
            print("[2] Title - "+title)
        if("DAY" in time):
            print("[3] Extra - DAY TIME")
        elif("NIGHT" in time):
            print("[3] - NIGHT TIME")
        else:
            print("[3] "+time)
        if(action == "0"):
            print("[4] Action - Not Set")
        else:
            print("[4] Action - "+action)
        print("[5] Next")

        option = int(input("Enter an Option number to edit : "))
        if(option == 1):
            ask_where()
        elif(option == 2):
            ask_title()
        elif(option == 3):
            ask_time()
        elif(option == 4):
            ask_action()
        elif(option == 5):
            save_title()
            #Goto character setting
            steps = 2
        else:
            print("invalid option")
    elif(steps == 2):

        print("Your character settings :- ")
        if(len(charac) == 0):
            print("[0] No characters are here - add")
            print("[666] Next Scene")
        else:
            print("[0] Total "+ str(len(charac))+" persons are here - add more")
            cou = 0
            for x in charac:
                cou = cou + 1
                print("["+str(cou)+"]")
                print("Name : "+ charac[x][0])
                print("Dialogue : "+ charac[x][1])
                print("paratheticy : "+ charac[x][2])
        option = int(input("Enter an Option number to add/edit : "))
        if(option <= len(charac) and option != 0):
            print("Edit this person ?")
            print("Name : "+ charac[option][0])
            option2 = int(input("'1' for yes , '2' for go back : "))
            if(option2 == 1):
                print("Leave blank for no-editing")
                name = ""
                dialogue = ""
                paratheticy = ""
                name = input("Re-enter name : ")
                name = name.upper()
                dialogue = input("Re-enter dialogue : ")
                paratheticy = input("Re-enter additional : ")
                paratheticy = "("+paratheticy+")"
                if(len(name) != 0):
                    charac[option][0] = name
                if(len(dialogue) != 0):
                    charac[option][1] = dialogue
                if(paratheticy != "()"):
                    charac[option][2] = paratheticy

        elif(option == 0):
            char_flag = 1
            loop_flag = 0

            while(char_flag):
                number = index
                name = ""
                if(loop_flag == 0):
                    name = input("Enter the name of person : ")
                    loop_flag = 1
                else:
                    ex = int(input("Enter '1' to add more , '2' to go back : "))
                    if(ex == 1):
                        name = input("Enter the name of next person : ")
                    else:
                        char_flag = 0
                        break
                name = name.upper()
                print("Leave blank and enter if none")
                dialogue = input("Enter his/her the dialogue : ")
                paratheticy = input("Any additional example 'Angry' : ")
                paratheticy = "(" + paratheticy +")"
                charac[index] = [name,dialogue,paratheticy]
                index = index + 1

            save_char(charac)
        elif(option == 666):
            count = count + 1
            break
        else:
            print("invalid option")
